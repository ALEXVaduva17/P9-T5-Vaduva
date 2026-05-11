"""
Script to generate SRS_Fitness_Center_v1.3_Corectat.docx
applying all corrections identified during the conformity audit.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, color=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_bullet(doc, text, bold=False, color=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

doc = Document()

# ─── Cover / Title ──────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Software Requirements Specification")
r.bold = True
r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("for Distributed Information System for Fitness Center Management")
r2.font.size = Pt(14)

doc.add_paragraph()
meta = [
    ("Version:", "1.3 approved"),
    ("Date:", "May 11, 2026"),
    ("Status:", "Approved"),
    ("Prepared by:", "Development Team"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(label + " ")
    rb.bold = True
    p.add_run(value)

doc.add_page_break()

# ─── Revision History ───────────────────────────────────────────────────────
add_heading(doc, "Revision History", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, h in enumerate(["Name", "Date", "Reason For Changes", "Version"]):
    hdr_cells[i].text = h
    hdr_cells[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr_cells[i], "D9D9D9")

revisions = [
    ("Dev Team", "March 09, 2026", "Initial version", "1.0"),
    ("Dev Team", "May 11, 2026", "Corrections: REQ-2 updated, duplicates removed, DB schema updated.", "1.1"),
    ("Dev Team", "May 11, 2026",
     "Corrections v1.2: Aligned with Tema_13. Reverted REQ-2 to 2 user types. "
     "Removed duplicate REQ-9. Deleted REQ-15 completely. Added missing requirement "
     "for viewing active subscription. Clarified REQ-6, REQ-7, REQ-8. DB schema updated.", "1.2"),
    ("Dev Team", "May 11, 2026",
     "Corrections v1.3: QA Audit — (1) Confirmed 2 user roles (Admin/Member; Trainer has NO login). "
     "(2) Added missing REQ-17 (Admin views all members' payment history; "
     "Members view only their own). (3) Corrected Section 4.2 description to remove erroneous "
     "'equipment inventory' from member-management wording. (4) Added explicit note that Trainers "
     "are NOT system users. (5) Verified DB schema covers all FRs. "
     "(6) Confirmed REQ-9 / REQ-15 removals are clean.", "1.3"),
]
for rev in revisions:
    row_cells = table.add_row().cells
    for i, val in enumerate(rev):
        row_cells[i].text = val

doc.add_paragraph()

# ─── Section 1 ──────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", level=1)
add_heading(doc, "1.1 Purpose", level=2)
add_paragraph(doc,
    "This document specifies the software requirements for the Distributed Information System "
    "for Fitness Center Management, version 1.3. It covers the entire system, including modules "
    "for authentication, member management, subscriptions, reservations, payments, and reporting.")

add_heading(doc, "1.2 Document Conventions", level=2)
add_paragraph(doc,
    "Requirements are identified by the prefix REQ- followed by a sequential number. "
    "Priorities are indicated as: H=High, M=Medium, L=Low. The term user refers generically "
    "to any person interacting with the system.")
add_paragraph(doc, "Formatting Legend:", bold=True)
p = doc.add_paragraph()
p.add_run("• Normal text ").bold = False
p.add_run("— unchanged requirement.\n")
r_new = p.add_run("• Green / [NEW] ")
r_new.font.color.rgb = RGBColor(0, 128, 0)
r_new.bold = True
p.add_run("— newly added requirement in this version.\n")
r_mod = p.add_run("• Blue / [MODIFIED] ")
r_mod.font.color.rgb = RGBColor(0, 70, 180)
r_mod.bold = True
p.add_run("— requirement modified in this version.\n")
r_del = p.add_run("• Red strikethrough / [REMOVED] ")
r_del.font.color.rgb = RGBColor(180, 0, 0)
r_del.bold = True
p.add_run("— requirement removed (referenced for traceability; no active text follows).")

add_heading(doc, "1.3 Intended Audience and Reading Suggestions", level=2)
add_paragraph(doc,
    "This document is intended for development teams, testers, project managers, and the end client. "
    "It is recommended to read sections 1 and 2 first for general context, followed by sections 3 and 4 "
    "for technical details.")

add_heading(doc, "1.4 Product Scope", level=2)
add_paragraph(doc,
    "The application enables complete management of activities within a fitness and wellness center. "
    "The system provides a web interface and/or a dedicated application through which members and "
    "administrators can access platform services. Key objectives include automating administrative "
    "processes, improving the member experience, and providing detailed reports to management.")

add_heading(doc, "1.5 References", level=2)
add_bullet(doc, "IEEE Std 830-1998 - Recommended Practice for Software Requirements Specifications")
add_bullet(doc, "Project Vision and Scope Document (version 1.0) / Tema 13")
add_bullet(doc, "Internal UI Style Guide")

# ─── Section 2 ──────────────────────────────────────────────────────────────
add_heading(doc, "2. Overall Description", level=1)
add_heading(doc, "2.1 Product Perspective", level=2)
add_paragraph(doc,
    "The system is a new, self-contained product composed of: a web client (browser) or dedicated "
    "mobile application, a database server storing information about members, subscriptions, and trainers, "
    "an application server, and a web server. Users access the service exclusively through the user interface.")

add_heading(doc, "2.2 Product Functions", level=2)
add_bullet(doc, "Secure authentication via username and password")
add_bullet(doc, "Viewing current subscription information")
add_bullet(doc, "Reserving spots in group training sessions")
add_bullet(doc, "Online payment of subscriptions and viewing payment history")
add_bullet(doc, "Managing trainer schedules and facility availability")

add_heading(doc, "2.3 User Classes and Characteristics", level=2)

# CORRECTION: Explicit note that Trainer is NOT a system user
p_note = doc.add_paragraph()
r_mod_label = p_note.add_run("[MODIFIED v1.3] ")
r_mod_label.font.color.rgb = RGBColor(0, 70, 180)
r_mod_label.bold = True
p_note.add_run(
    "The system architecture defines exactly TWO (2) roles that can authenticate in the system, "
    "as specified in Tema 13. Trainers are internal entities managed by Administrators for scheduling "
    "purposes and do NOT have system login credentials."
)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(["User Class", "Characteristics", "Privileges"]):
    hdr2[i].text = h
    hdr2[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr2[i], "D9D9D9")

user_classes = [
    ("Administrator",
     "Technical or administrative staff with full system access",
     "Add/edit members, modify subscription prices, manage equipment inventory, "
     "manage trainer schedules, view all members' payment history, view reports."),
    ("Member",
     "Fitness center client",
     "View own active subscription, make class reservations, access exclusively "
     "personal payment history."),
    # CORRECTION v1.3: Trainer row clarified as non-user entity
    ("Trainer (non-user entity)",
     "Fitness center employee who leads classes; managed by Administrators",
     "NO system login. Administrators manage trainer records on their behalf."),
]
for uc in user_classes:
    row_cells = table2.add_row().cells
    for i, val in enumerate(uc):
        row_cells[i].text = val
# Grey-out the Trainer row to emphasize it is not a login role
for cell in table2.rows[-1].cells:
    set_cell_bg(cell, "F2F2F2")
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(120, 120, 120)
        run.italic = True

doc.add_paragraph()

add_heading(doc, "2.4 Operating Environment", level=2)
add_paragraph(doc,
    "The system will operate as a web application accessible from any modern browser (Chrome, Firefox, "
    "Safari, Edge). The application server will run on a Linux platform. The database will be hosted on "
    "a dedicated server with automatic backup support.")

add_heading(doc, "2.5 Design and Implementation Constraints", level=2)
add_bullet(doc, "The system must comply with GDPR regulations for the protection of personal data")
add_bullet(doc, "All communications will be secured via HTTPS/TLS")
add_bullet(doc, "Online payments will be processed through a PCI-DSS certified payment gateway")

add_heading(doc, "2.6 User Documentation", level=2)
add_bullet(doc, "Online user manual (HTML)")
add_bullet(doc, "Quick start guide (PDF)")
add_bullet(doc, "Contextual help system integrated into the application")

add_heading(doc, "2.7 Assumptions and Dependencies", level=2)
add_bullet(doc, "A functional online payment gateway is assumed to be available")
add_bullet(doc, "Clients have internet access and a modern browser")
add_bullet(doc, "The system depends on an external email/SMS service for notifications")

# ─── Section 3 ──────────────────────────────────────────────────────────────
add_heading(doc, "3. External Interface Requirements", level=1)

add_heading(doc, "3.1 User Interfaces", level=2)
add_paragraph(doc,
    "The web interface will be responsive, compatible with mobile and desktop devices, and will follow "
    "WCAG 2.1 accessibility principles. Every screen will include a navigation menu, clear error messages, "
    "and confirmation dialogs for irreversible actions.")

add_heading(doc, "3.2 Hardware Interfaces", level=2)
add_paragraph(doc,
    "The system requires no special hardware on the user side. Servers will require a minimum of 8 GB RAM, "
    "SSD storage, and a redundant network connection.")

add_heading(doc, "3.3 Software Interfaces", level=2)
add_paragraph(doc,
    "The system will integrate with: a payment gateway (e.g., Stripe or PayU) for processing transactions, "
    "an SMTP service for sending notification emails, and optionally an external calendar system for "
    "schedule synchronization.")

add_heading(doc, "3.4 Communications Interfaces", level=2)
add_paragraph(doc,
    "All client-server communications will use the HTTPS protocol (TLS 1.2 or higher). Internal APIs "
    "will follow REST architecture. Notifications will be sent via email and/or SMS.")

# ─── Section 4 ──────────────────────────────────────────────────────────────
add_heading(doc, "4. System Features", level=1)

# 4.1 Authentication
add_heading(doc, "4.1 Authentication and Access Control", level=2)
add_paragraph(doc, "Description and Priority: High (H)", bold=True)
add_paragraph(doc,
    "The system must provide a secure authentication mechanism and control access based on the user's role.")
add_bullet(doc, "REQ-1: Access to the application is protected by a username and password.")
p_req2 = doc.add_paragraph(style='List Bullet')
r_tag = p_req2.add_run("[MODIFIED v1.2] ")
r_tag.font.color.rgb = RGBColor(0, 70, 180)
r_tag.bold = True
p_req2.add_run(
    "REQ-2: There are exactly TWO (2) user types that can authenticate in the system: "
    "Administrators and Members. Trainers are NOT system users and do not have login credentials."
)
add_bullet(doc, "REQ-3: The system will block access after 5 consecutive failed login attempts.")

# 4.2 Member Management
add_heading(doc, "4.2 Member Management", level=2)
add_paragraph(doc, "Description and Priority: High (H)", bold=True)

# CORRECTION v1.3: Removed 'equipment inventory' from description (it's an Admin capability, not member mgmt description error)
p_42desc = doc.add_paragraph()
r_42tag = p_42desc.add_run("[MODIFIED v1.3] ")
r_42tag.font.color.rgb = RGBColor(0, 70, 180)
r_42tag.bold = True
p_42desc.add_run(
    "Administrators can fully manage the member database, including adding, editing members "
    "and managing member-related data. Equipment inventory is managed separately under Admin capabilities."
)

add_bullet(doc, "REQ-4: Administrators can add/edit members, modify subscription prices, and manage equipment inventory.")
add_bullet(doc, "REQ-5: Members are identified by: first name, last name, phone number, and email address.")

p_req6 = doc.add_paragraph(style='List Bullet')
r6_tag = p_req6.add_run("[MODIFIED v1.2] ")
r6_tag.font.color.rgb = RGBColor(0, 70, 180)
r6_tag.bold = True
p_req6.add_run(
    "REQ-6: When a subscription expires, the member is automatically moved to the restricted_clients table. "
    "A restricted client is functionally blocked from making new class reservations. The member record "
    "remains in the members table; only a reference is added to restricted_clients."
)

# 4.3 Subscription and Payment Management
add_heading(doc, "4.3 Subscription and Payment Management", level=2)
add_paragraph(doc, "Description and Priority: High (H)", bold=True)
add_paragraph(doc,
    "This module allows viewing of active subscriptions, online payment, viewing payment history, "
    "and calculation of the amount due for personalized subscriptions.")

p_req16 = doc.add_paragraph(style='List Bullet')
r16_tag = p_req16.add_run("[NEW v1.2] ")
r16_tag.font.color.rgb = RGBColor(0, 128, 0)
r16_tag.bold = True
p_req16.add_run("REQ-16: Members can view detailed information regarding their current active subscription.")

p_req7 = doc.add_paragraph(style='List Bullet')
r7_tag = p_req7.add_run("[MODIFIED v1.2] ")
r7_tag.font.color.rgb = RGBColor(0, 70, 180)
r7_tag.bold = True
p_req7.add_run(
    "REQ-7: Members can pay for subscriptions online and view EXCLUSIVELY their personal payment history. "
    "Members do NOT have access to other members' payment records."
)

# NEW: REQ-17 — Admin payment history view
p_req17 = doc.add_paragraph(style='List Bullet')
r17_tag = p_req17.add_run("[NEW v1.3] ")
r17_tag.font.color.rgb = RGBColor(0, 128, 0)
r17_tag.bold = True
p_req17.add_run(
    "REQ-17: Administrators can view the full payment history of all members for auditing and reporting purposes."
)

p_req8 = doc.add_paragraph(style='List Bullet')
r8_tag = p_req8.add_run("[MODIFIED v1.2] ")
r8_tag.font.color.rgb = RGBColor(0, 70, 180)
r8_tag.bold = True
p_req8.add_run(
    "REQ-8: A personalized subscription can be initiated by either a Member or an Administrator. "
    "The payment amount is calculated as: Amount = Base Fee + (No. of personal training sessions × 50 RON)."
)

add_note(doc, "(Note: REQ-9 was removed in v1.2 — it was a confirmed duplicate of REQ-13. No active text.)")

# 4.4 Training Class Reservations
add_heading(doc, "4.4 Training Class Reservations", level=2)
add_paragraph(doc, "Description and Priority: Medium (M)", bold=True)
add_paragraph(doc, "Members can reserve spots in available group training sessions.")
add_bullet(doc,
    "REQ-10: A member can have only one active subscription but may make multiple reservations "
    "for different classes.")
add_bullet(doc,
    "REQ-11: For each reservation, the system stores the date, time, and facility where "
    "the activity takes place.")

# 4.5 Trainer and Facility Management
add_heading(doc, "4.5 Trainer and Facility Management", level=2)
add_paragraph(doc, "Description and Priority: Medium (M)", bold=True)
add_paragraph(doc,
    "This module enables management of trainer schedules and training facility availability. "
    "All trainer-related records are managed exclusively by Administrators.")
add_bullet(doc, "REQ-12: Administrators can define each trainer's schedule and available facilities.")

# 4.6 Reporting
add_heading(doc, "4.6 Reporting", level=2)
add_paragraph(doc, "Description and Priority: Medium (M)", bold=True)
add_paragraph(doc, "Reporting modules for management oversight.")
add_bullet(doc, "REQ-13: A report of total amounts collected from members for a specific month can be generated.")
add_bullet(doc, "REQ-14: A report showing the occupancy rate of training facilities can be generated.")
add_note(doc, "(Note: Former duplicate REQ-15 has been fully removed in v1.2. No active text.)")

# ─── Section 5 ──────────────────────────────────────────────────────────────
add_heading(doc, "5. Other Nonfunctional Requirements", level=1)

add_heading(doc, "5.1 Performance Requirements", level=2)
add_bullet(doc, "Server response time for common operations shall not exceed 2 seconds under normal load.")
add_bullet(doc, "The system shall support at least 100 concurrent users without performance degradation.")
add_bullet(doc, "Monthly reports shall be generated within 10 seconds.")

add_heading(doc, "5.2 Safety Requirements", level=2)
add_bullet(doc, "Member data shall be backed up automatically every day.")
add_bullet(doc, "In the event of a critical failure, the system shall recover to the last stable state within 4 hours.")

add_heading(doc, "5.3 Security Requirements", level=2)
add_bullet(doc, "Passwords shall be stored exclusively as hashes (bcrypt or equivalent).")
add_bullet(doc, "Sessions shall expire after 30 minutes of inactivity.")
add_bullet(doc, "All transmitted data shall be encrypted via TLS 1.2 or higher.")
add_bullet(doc, "The system shall comply with GDPR provisions for personal data.")

add_heading(doc, "5.4 Software Quality Attributes", level=2)
add_bullet(doc, "Availability: The system shall be available 99.5% of the time, excluding planned maintenance.")
add_bullet(doc, "Maintainability: Source code shall follow documented coding standards.")
add_bullet(doc, "Portability: The web application shall function on the last 2 major versions of mainstream browsers.")
add_bullet(doc, "Usability: A new user shall be able to complete a reservation without prior training.")

add_heading(doc, "5.5 Business Rules", level=2)
add_bullet(doc, "A member with an expired subscription cannot make new reservations.")
add_bullet(doc, "Subscription prices can only be modified by an Administrator.")
add_bullet(doc, "A personalized subscription includes the base fee plus 50 RON per personal training session.")
add_bullet(doc, "Members can view only their own payment history; Administrators can view all payment histories.")

# ─── Section 6 ──────────────────────────────────────────────────────────────
add_heading(doc, "6. Other Requirements", level=1)
add_paragraph(doc,
    "The system will support Romanian as the primary interface language. Date formats will follow "
    "the European standard (DD.MM.YYYY). The currency used for all transactions will be RON (Romanian Leu).")

# ─── Section 7 ──────────────────────────────────────────────────────────────
add_heading(doc, "7. Technical Architecture", level=1)

add_heading(doc, "7.1 Overview", level=2)
add_paragraph(doc,
    "The system follows a standard client-server architecture with a decoupled frontend and backend. "
    "The backend exposes a REST API consumed by the frontend. All data is persisted in a relational database.")

add_heading(doc, "7.2 Backend & 7.3 Frontend", level=2)
add_paragraph(doc,
    "The backend is built with FastAPI (Python), PostgreSQL, SQLAlchemy, and Alembic. "
    "The frontend is a React single-page application communicating exclusively via HTTP/HTTPS.")

add_heading(doc, "7.4 Database Schema (Core Entities)", level=2)

p_schema_note = doc.add_paragraph()
r_sn_tag = p_schema_note.add_run("[MODIFIED v1.3] ")
r_sn_tag.font.color.rgb = RGBColor(0, 70, 180)
r_sn_tag.bold = True
p_schema_note.add_run(
    "Schema reflects exactly 2 authenticated user roles (admin/member). Trainers are managed entities "
    "without system login. The payments table includes member_id to support REQ-17 (Admin views all payments)."
)

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(["Table", "Key Columns"]):
    hdr3[i].text = h
    hdr3[i].paragraphs[0].runs[0].bold = True
    set_cell_bg(hdr3[i], "D9D9D9")

schema_rows = [
    ("users", "email, hashed_password, role (admin | member), is_active"),
    ("members", "user_id (FK → users), first_name, last_name, phone, subscription_status"),
    ("trainers",
     "first_name, last_name, phone  ← NO user_id FK; trainers are NOT system users"),
    ("subscriptions",
     "member_id (FK → members), type, base_fee, pt_sessions, total_amount, start_date, end_date, status"),
    ("payments",
     "subscription_id (FK → subscriptions), member_id (FK → members), amount, currency, "
     "gateway_reference, paid_at  ← member_id added to support REQ-17"),
    ("classes", "trainer_id (FK → trainers), facility_id (FK → facilities), name, start_time, end_time, capacity"),
    ("reservations", "member_id (FK → members), class_id (FK → classes), reserved_at, status"),
    ("facilities", "name, capacity, description"),
    ("equipment", "name, category, quantity, condition, facility_id (FK → facilities), notes"),
    ("restricted_clients",
     "member_id (FK → members), subscription_id (FK → subscriptions), restricted_at, reason"),
]
for tbl_name, cols in schema_rows:
    row_cells = table3.add_row().cells
    row_cells[0].text = tbl_name
    row_cells[1].text = cols
    if "NO user_id" in cols or "← member_id added" in cols:
        for cell in row_cells:
            set_cell_bg(cell, "E8F5E9")

doc.add_paragraph()

# ─── Appendices ─────────────────────────────────────────────────────────────
add_heading(doc, "Appendices", level=1)

add_heading(doc, "Appendix A: Glossary", level=2)
glossary = [
    ("Member", "A natural person registered in the system with an active or expired subscription."),
    ("Administrator", "A user with full system management rights."),
    ("Trainer",
     "A fitness center employee who leads group classes and personal training sessions. "
     "Trainers are NOT system users and have NO login credentials."),
    ("Restricted Client",
     "A member whose subscription has expired and who has been automatically transferred to the "
     "restricted_clients table, preventing new reservations."),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    r_term = p.add_run(term + ": ")
    r_term.bold = True
    p.add_run(definition)
    p.paragraph_format.left_indent = Inches(0.3)

add_heading(doc, "Appendix B: Analysis Models", level=2)
add_paragraph(doc,
    "This appendix will include relevant analysis diagrams during the detailed design phase.")

add_heading(doc, "Appendix C: To Be Determined List", level=2)
add_paragraph(doc,
    "No TBD items at this time. This list will be updated as the project progresses.")

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = "SRS_Fitness_Center_v1.3_Corectat.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
