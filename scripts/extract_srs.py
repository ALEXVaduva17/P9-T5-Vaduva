from docx import Document

doc = Document('SRS_Fitness_Center_v1.2_Complet.docx')
full_text = []

for para in doc.paragraphs:
    if para.text.strip():
        full_text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
        if row_text:
            full_text.append('[TABLE ROW] ' + row_text)

with open('srs_extracted.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(full_text))

print(f'SRS extracted: {len(full_text)} lines')
